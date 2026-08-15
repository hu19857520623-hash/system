# -*- coding: utf-8 -*-
"""将 docs 目录下 HTML 设计文档转换为 Word (.docx) 并保存到桌面。"""
import base64
import io
import re
import sys
from pathlib import Path

try:
    from bs4 import BeautifulSoup, NavigableString
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', 'beautifulsoup4', 'lxml', 'requests', '-q'])
    from bs4 import BeautifulSoup, NavigableString
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

import requests

DOCS_DIR = Path(__file__).parent
DESKTOP = Path.home() / 'Desktop'
HTML_FILES = [
    ('OMS系统方案设计图.html', 'OMS系统方案设计图.docx'),
    ('OMS全流程图.html', 'OMS全流程图.docx'),
    ('客户流程与人员介入图.html', '客户流程与人员介入图.docx'),
]

MERMAID_TIMEOUT = 30


def set_doc_font(doc: Document):
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text.strip(), level=min(level, 3))
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_para(doc: Document, text: str, bold: bool = False, color=None):
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def strip_html_text(el) -> str:
    if el is None:
        return ''
    text = el.get_text(separator=' ', strip=True)
    text = re.sub(r'\s+', ' ', text)
    return text


def mermaid_to_image(code: str) -> bytes | None:
    code = code.strip()
    if not code:
        return None
    encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png&bgColor=white'
    try:
        r = requests.get(url, timeout=MERMAID_TIMEOUT)
        if r.status_code == 200 and r.headers.get('content-type', '').startswith('image'):
            return r.content
    except Exception:
        pass
    return None


def add_mermaid_block(doc: Document, code: str):
    img = mermaid_to_image(code)
    if img:
        try:
            doc.add_picture(io.BytesIO(img), width=Inches(6.2))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        except Exception:
            pass
    add_para(doc, '【流程图源码 — 可在支持 Mermaid 的工具中渲染】', bold=True)
    p = doc.add_paragraph()
    run = p.add_run(code.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(8)


def html_table_to_docx(doc: Document, table_el):
    rows = table_el.find_all('tr')
    if not rows:
        return
    cols = max(len(r.find_all(['th', 'td'])) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < cols:
                tbl.rows[i].cells[j].text = strip_html_text(cell)
                if cell.name == 'th':
                    for p in tbl.rows[i].cells[j].paragraphs:
                        for run in p.runs:
                            run.bold = True


def convert_oms_blueprint(doc: Document, soup: BeautifulSoup) -> bool:
    """解析 OMS 方案设计图（.sheet 布局）。"""
    sheets = soup.find_all(class_='sheet')
    if not sheets:
        page = soup.find(class_='canvas') or soup.find(id='diagram')
        if not page:
            return False
        sheets = [page]

    for si, sheet in enumerate(sheets):
        if si > 0:
            doc.add_page_break()

        title = sheet.find(class_='banner') or sheet.find(class_='top-banner')
        if title:
            h = title.find('h1') if title.name != 'h1' else title
            if not h and title.name == 'div':
                h = title.find('h1')
            if h:
                add_heading(doc, strip_html_text(h), 1)
            sub = title.find('p') if title.name != 'p' else None
            if sub:
                add_para(doc, strip_html_text(sub))

        # 角色
        roles = sheet.find_all(class_='role-item') or sheet.find_all(class_='role')
        if roles and not sheet.find(class_='perm-role'):
            add_heading(doc, '使用角色', 2)
            for r in roles:
                name = r.find(class_='role-name') or r.find(class_='role-t')
                sub = r.find(class_='role-sub') or r.find(class_='role-d')
                add_para(doc, f"• {strip_html_text(name)}：{strip_html_text(sub)}")

        # 五大中心
        pillars = sheet.find_all(class_='pillar')
        cols = sheet.find_all(class_='col')
        modules = pillars or cols
        if modules:
            add_heading(doc, '功能模块', 2)
            tbl = doc.add_table(rows=len(modules) + 1, cols=2)
            tbl.style = 'Table Grid'
            tbl.rows[0].cells[0].text = '中心'
            tbl.rows[0].cells[1].text = '功能点'
            for i, p in enumerate(modules):
                head = p.find(class_='pillar-title') or p.find(class_='col-h')
                feats = p.find_all(class_='feat') or p.find_all(class_='cell')
                tbl.rows[i + 1].cells[0].text = strip_html_text(head)
                tbl.rows[i + 1].cells[1].text = '、'.join(
                    strip_html_text(f.find('span')) for f in feats if f.find('span')
                )

        # 对外输出
        out_cards = sheet.find_all(class_='out-card') or sheet.find_all(class_='out-item')
        if out_cards:
            add_heading(doc, 'OMS 对外输出', 2)
            for c in out_cards:
                add_para(doc, '• ' + strip_html_text(c))

        # 核心流程
        flow_lbls = sheet.find_all(class_='flow-lbl')
        steps = sheet.find_all(class_='step')
        flow_text = flow_lbls or steps
        if flow_text:
            add_heading(doc, 'OMS 核心流程', 2)
            add_para(doc, ' → '.join(strip_html_text(x) for x in flow_text))

        # 基础支撑
        infra_nodes = sheet.find_all(class_='infra-node') or sheet.find_all(class_='infra-cell')
        if infra_nodes:
            add_heading(doc, '基础支撑层', 2)
            labels = []
            for x in infra_nodes:
                span = x.find('span')
                labels.append(strip_html_text(span) if span else strip_html_text(x))
            add_para(doc, ' · '.join(labels))

        # 权限角色（新版）
        perm_roles = sheet.find_all(class_='perm-role')
        if perm_roles:
            add_heading(doc, '权限角色说明', 2)
            for card in perm_roles:
                h = card.find('h3')
                add_para(doc, strip_html_text(h), bold=True)
                for li in card.find_all('li'):
                    add_para(doc, '  - ' + strip_html_text(li))

        # 权限卡片（旧版）
        perm_cards = sheet.find_all(class_='perm-card')
        if perm_cards:
            add_heading(doc, '权限角色说明', 2)
            for card in perm_cards:
                h = card.find(class_='perm-card-h')
                add_para(doc, strip_html_text(h), bold=True)
                for li in card.find_all('li'):
                    add_para(doc, '  - ' + strip_html_text(li))

        # 权限分配流程
        pf_steps = sheet.find_all(class_='pf-step')
        if pf_steps:
            add_heading(doc, '权限分配流程', 2)
            parts = []
            for s in pf_steps:
                b = s.find('b')
                sm = s.find('small')
                parts.append(f"{strip_html_text(b)}（{strip_html_text(sm)}）" if sm else strip_html_text(b))
            add_para(doc, ' → '.join(parts))

        # 客户业务场景
        biz_tracks = sheet.find_all(class_='biz-track')
        if biz_tracks:
            add_heading(doc, '客户业务场景', 2)
            for track in biz_tracks:
                h = track.find('h4')
                steps = track.find_all(class_='biz-step')
                note = track.find(class_='biz-note')
                flow = ' → '.join(strip_html_text(s) for s in steps)
                add_para(doc, f"{strip_html_text(h)}：{flow}", bold=True)
                if note:
                    add_para(doc, strip_html_text(note))

        # 发货来源 & 库存来源
        meta_panels = sheet.find_all(class_='meta-panel')
        if meta_panels:
            for panel in meta_panels:
                h = panel.find('h4')
                tags = panel.find_all(class_='meta-tag')
                add_heading(doc, strip_html_text(h), 2)
                add_para(doc, ' · '.join(strip_html_text(t) for t in tags))

        # 权限矩阵
        matrix = sheet.find('table.matrix')
        if matrix:
            add_heading(doc, '模块权限矩阵', 2)
            html_table_to_docx(doc, matrix)

    return True


def convert_html_to_docx(html_path: Path, out_path: Path):
    soup = BeautifulSoup(html_path.read_text(encoding='utf-8'), 'lxml')
    doc = Document()
    set_doc_font(doc)

    if convert_oms_blueprint(doc, soup):
        pass
    else:
        # 旧版 / 通用 HTML
        h1 = soup.find('h1')
        if h1:
            add_heading(doc, strip_html_text(h1), 1)
        header_p = soup.find('header')
        if header_p:
            for p in header_p.find_all('p', recursive=False):
                add_para(doc, strip_html_text(p))
            meta = header_p.find(class_='meta')
            if meta:
                add_para(doc, strip_html_text(meta))
        doc.add_paragraph()

    for section in soup.find_all('section'):
        h2 = section.find('h2')
        if h2:
            add_heading(doc, strip_html_text(h2), 2)

        desc = section.find(class_='desc')
        if desc:
            add_para(doc, strip_html_text(desc))

        callout = section.find(class_='callout')
        if callout:
            add_para(doc, strip_html_text(callout), color=RGBColor(0x92, 0x40, 0x0E))

        for child in section.children:
            if not hasattr(child, 'name'):
                continue
            if child.name == 'table':
                html_table_to_docx(doc, child)
                doc.add_paragraph()
            elif child.name == 'div' and 'mermaid' in (child.get('class') or []):
                code = child.get_text()
                add_mermaid_block(doc, code)
                doc.add_paragraph()
            elif child.name == 'div' and child.find('table'):
                for t in child.find_all('table'):
                    html_table_to_docx(doc, t)
                    doc.add_paragraph()

        legend = section.find(class_='legend')
        if legend:
            add_para(doc, '图例：' + strip_html_text(legend))

        doc.add_paragraph()

    footer = soup.find('footer')
    if footer:
        add_para(doc, strip_html_text(footer), color=RGBColor(0x64, 0x74, 0x8B))

    doc.save(str(out_path))
    print(f'OK: {out_path}')


def main():
    DESKTOP.mkdir(parents=True, exist_ok=True)
    for html_name, docx_name in HTML_FILES:
        html_path = DOCS_DIR / html_name
        if not html_path.exists():
            print(f'SKIP (not found): {html_path}')
            continue
        out_path = DESKTOP / docx_name
        print(f'Converting {html_name} -> {out_path}')
        convert_html_to_docx(html_path, out_path)
    print(f'\nAll Word files saved to: {DESKTOP}')


if __name__ == '__main__':
    main()
