# -*- coding: utf-8 -*-
"""Generate assistant probation assessment Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "黑体", 18, True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p.paragraph_format.space_after = Pt(18)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 14, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, "黑体", 12, True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_blank_line(doc, text=""):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, "宋体", 12, False)
    p.paragraph_format.space_after = Pt(4)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, "宋体", 10, row_idx == 0)
                paragraph.paragraph_format.space_after = Pt(2)
            if row_idx == 0:
                set_cell_shading(cell, "D9E2F3")


def create_document(output_path: str):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    add_title(doc, "助理试用期考核表")
    add_subtitle(doc, "（系统数据维护 & 客户培训）")

    add_heading(doc, "基本信息")
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("姓名", ""),
        ("部门", ""),
        ("岗位", "系统数据维护 / 客户培训助理"),
        ("入职日期", "    年    月    日"),
        ("考核日期", "    年    月    日"),
        ("试用期", "1 个月"),
        ("考核人", ""),
        ("被考核人", ""),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    style_table(info_table)
    doc.add_paragraph()

    add_heading(doc, "一、考核说明")
    add_heading(doc, "1.1 岗位职责", level=2)
    for item in [
        "负责 ERP 系统基础数据的维护（商品、供应商、客户、入库单等）；",
        "对接新客户：系统开户、充值配置、基础操作培训；",
        "处理客户日常使用问题，无法独立解决时及时上报。",
    ]:
        add_body(doc, item, indent=True)

    add_heading(doc, "1.2 考核维度及权重", level=2)
    weight_table = doc.add_table(rows=5, cols=2)
    weight_table.style = "Table Grid"
    weight_data = [
        ("考核维度", "权重"),
        ("系统数据维护", "40%"),
        ("客户培训与支持", "35%"),
        ("工作态度与协作", "15%"),
        ("学习与成长", "10%"),
    ]
    for i, (a, b) in enumerate(weight_data):
        weight_table.rows[i].cells[0].text = a
        weight_table.rows[i].cells[1].text = b
    style_table(weight_table)
    doc.add_paragraph()

    add_heading(doc, "1.3 评分标准", level=2)
    score_table = doc.add_table(rows=5, cols=2)
    score_table.style = "Table Grid"
    score_data = [
        ("等级", "说明"),
        ("优秀", "超出岗位要求，可独立承担全部工作"),
        ("良好", "达到岗位要求，偶需指导"),
        ("合格", "基本达到要求，转正后需继续带教"),
        ("不合格", "未达到岗位要求"),
    ]
    for i, (a, b) in enumerate(score_data):
        score_table.rows[i].cells[0].text = a
        score_table.rows[i].cells[1].text = b
    style_table(score_table)
    doc.add_paragraph()

    add_heading(doc, "1.4 转正参考指标", level=2)
    ref_table = doc.add_table(rows=6, cols=2)
    ref_table.style = "Table Grid"
    ref_data = [
        ("指标", "合格参考线"),
        ("数据维护任务独立完成率", "≥ 80%"),
        ("重大数据错误次数", "≤ 2 次，且同类错误不重复"),
        ("独立完成客户 onboarding", "≥ 1–2 个"),
        ("工作时间内客户消息响应", "≤ 30 分钟"),
        ("客户明确投诉", "0 次"),
    ]
    for i, (a, b) in enumerate(ref_data):
        ref_table.rows[i].cells[0].text = a
        ref_table.rows[i].cells[1].text = b
    style_table(ref_table)
    doc.add_paragraph()

    add_heading(doc, "二、被考核人自评")
    add_body(doc, "（请被考核人填写，300 字左右）")
    for label in [
        "2.1 试用期主要工作内容：",
        "2.2 工作成果与亮点：",
        "2.3 存在的不足及改进计划：",
        "2.4 对岗位和团队的建议：",
    ]:
        add_heading(doc, label, level=2)
        for _ in range(3):
            add_blank_line(doc, "_______________________________________________________________________________")
    add_body(doc, "被考核人签字：____________　　　　日期：____年____月____日")
    doc.add_paragraph()

    add_heading(doc, "三、考核人评价")

    add_heading(doc, "3.1 系统数据维护（权重 40%）", level=2)
    maint_table = doc.add_table(rows=5, cols=6)
    maint_table.style = "Table Grid"
    maint_headers = ["考核项", "优秀", "良好", "合格", "不合格", "评价"]
    maint_rows = [
        "操作熟练度（各模块独立操作能力）",
        "数据准确性（录入、导入错误率）",
        "规范意识（是否按流程操作）",
        "异常处理（查日志、重试、上报）",
    ]
    for j, h in enumerate(maint_headers):
        maint_table.rows[0].cells[j].text = h
    for i, row_text in enumerate(maint_rows, start=1):
        maint_table.rows[i].cells[0].text = row_text
        for j in range(1, 5):
            maint_table.rows[i].cells[j].text = "□"
        maint_table.rows[i].cells[5].text = ""
    style_table(maint_table)
    doc.add_paragraph()

    add_body(doc, "具体事例（必填）：")
    add_body(doc, "做得好的：")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_body(doc, "出现过的问题：")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_body(doc, "本项综合评定：  □ 优秀    □ 良好    □ 合格    □ 不合格")
    doc.add_paragraph()

    add_heading(doc, "3.2 客户培训与支持（权重 35%）", level=2)
    train_table = doc.add_table(rows=5, cols=6)
    train_table.style = "Table Grid"
    train_rows = [
        "培训效果（客户能否独立操作）",
        "讲解能力（条理、耐心、表达）",
        "响应及时性（客户消息回复速度）",
        "客户反馈（投诉/好评）",
    ]
    for j, h in enumerate(maint_headers):
        train_table.rows[0].cells[j].text = h
    for i, row_text in enumerate(train_rows, start=1):
        train_table.rows[i].cells[0].text = row_text
        for j in range(1, 5):
            train_table.rows[i].cells[j].text = "□"
        train_table.rows[i].cells[5].text = ""
    style_table(train_table)
    doc.add_paragraph()

    add_body(doc, "培训客户数量：______ 个")
    add_body(doc, "客户/同事反馈：")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_body(doc, "本项综合评定：  □ 优秀    □ 良好    □ 合格    □ 不合格")
    doc.add_paragraph()

    add_heading(doc, "3.3 工作态度与协作（权重 15%）", level=2)
    attitude_table = doc.add_table(rows=5, cols=5)
    attitude_table.style = "Table Grid"
    attitude_headers = ["考核项", "优秀", "良好", "合格", "不合格"]
    attitude_rows = [
        "主动性（主动跟进、提醒、补位）",
        "责任心（出错承认、及时改正）",
        "沟通协作（对内对外清晰及时）",
        "保密意识（客户及业务数据）",
    ]
    for j, h in enumerate(attitude_headers):
        attitude_table.rows[0].cells[j].text = h
    for i, row_text in enumerate(attitude_rows, start=1):
        attitude_table.rows[i].cells[0].text = row_text
        for j in range(1, 5):
            attitude_table.rows[i].cells[j].text = "□"
    style_table(attitude_table)
    doc.add_paragraph()
    add_body(doc, "本项综合评定：  □ 优秀    □ 良好    □ 合格    □ 不合格")
    doc.add_paragraph()

    add_heading(doc, "3.4 学习与成长（权重 10%）", level=2)
    add_body(doc, "1 个月内新掌握的模块/流程：")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_body(doc, "是否仍需手把手带教：  □ 很少    □ 偶尔    □ 经常")
    add_body(doc, "转正后能否承担更多工作：  □ 能    □ 需再带一段时间    □ 不能")
    add_body(doc, "本项综合评定：  □ 优秀    □ 良好    □ 合格    □ 不合格")
    doc.add_paragraph()

    add_heading(doc, "四、综合结论")
    add_heading(doc, "4.1 考核结论（请选择一项）", level=2)
    add_body(doc, "□ 同意转正 — 达到岗位要求，可独立承担系统维护及客户培训工作")
    add_body(doc, "□ 延长试用期 2 周 — 基本达标但熟练度/培训能力尚需加强，延长原因：")
    add_blank_line(doc, "_______________________________________________________________________________")
    add_body(doc, "□ 不予转正 — 未达到岗位要求，主要原因：")
    add_blank_line(doc, "_______________________________________________________________________________")

    add_heading(doc, "4.2 综合评价（3–5 句话）", level=2)
    for _ in range(3):
        add_blank_line(doc, "_______________________________________________________________________________")

    add_heading(doc, "4.3 转正后期望（若同意转正）", level=2)
    for _ in range(2):
        add_blank_line(doc, "_______________________________________________________________________________")
    doc.add_paragraph()

    add_heading(doc, "五、审批签字")
    sign_table = doc.add_table(rows=5, cols=3)
    sign_table.style = "Table Grid"
    sign_data = [
        ("角色", "签字", "日期"),
        ("考核人（直属上级）", "", "    年    月    日"),
        ("被考核人", "", "    年    月    日"),
        ("部门负责人（如有）", "", "    年    月    日"),
        ("HR（如有）", "", "    年    月    日"),
    ]
    for i, row in enumerate(sign_data):
        for j, val in enumerate(row):
            sign_table.rows[i].cells[j].text = val
    style_table(sign_table)
    doc.add_paragraph()

    add_heading(doc, "附件：试用期工作记录表（考核人日常记录用）")
    log_table = doc.add_table(rows=11, cols=4)
    log_table.style = "Table Grid"
    log_headers = ["日期", "工作事项", "结果", "备注"]
    for j, h in enumerate(log_headers):
        log_table.rows[0].cells[j].text = h
    for i in range(1, 11):
        for j in range(4):
            log_table.rows[i].cells[j].text = ""
    style_table(log_table)
    add_body(doc, "记录说明：考核人可在试用期内随时填写，转正考核时作为评价依据。")
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("文档版本：V1.0  |  适用岗位：系统数据维护 & 客户培训助理  |  试用期：1 个月")
    set_run_font(run, "宋体", 9, False)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(output_path)
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    create_document(r"c:\Users\15693\Desktop\助理试用期考核表.docx")
