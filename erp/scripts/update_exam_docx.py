# -*- coding: utf-8 -*-
"""Update exam questions — replace some with ERP/WMS system operation questions."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
import os

SRC_CANDIDATES = [
    r"c:\Users\15693\Desktop\erp\docs\考题考核一期（修改版）.docx",
    r"d:\wx\xwechat_files\wxid_a0xp3budymkt12_1059\msg\file\2026-07\考题考核一期.docx",
]
OUT = r"c:\Users\15693\Desktop\erp\docs\考题考核一期（修改版）.docx"
OUT_ORIGINAL = r"d:\wx\xwechat_files\wxid_a0xp3budymkt12_1059\msg\file\2026-07\考题考核一期.docx"

LINE = "________________________________________________________________________"

# row_index: (module, seq, question)
QUESTIONS = {
    # ── 店铺类型（保留业务题 + 换入 ERP 操作题）──
    1: ("店铺类型", "1", f"本土店和跨境店，请至少列出 3 个主要区别（可从注册资质、回款方式、发货模式等角度说明）：{LINE}"),
    2: ("", "2", f"直邮模式下，货物通常通过什么物流方式从国内发往南非？{LINE}"),
    3: ("", "3", f"本土店与跨境店在回款周期和到账方式上有哪些主要差异？{LINE}"),
    4: ("", "4", f"产品链接上架后，长期不备货入平台仓会产生什么后果？{LINE}"),
    5: ("", "5", f"【系统操作】在 ERP「创建入库单」页面，SKU 明细是从哪里选取的？{LINE}"),
    6: ("", "6", f"【系统操作】创建入库单时勾选「自动推送 WMS」，若推送失败，系统状态显示什么？去哪里重试？{LINE}"),
    7: ("", "7", f"【系统操作】同一笔入库单中，SKU 能否来自不同采购单？系统如何限制？{LINE}"),
    8: ("", "8", f"【系统操作】创建入库单时，每条 SKU 明细除数量外还必须填写哪些字段？{LINE}"),
    9: ("", "9", f"店铺出现哪些严重违规情况可能导致封店或限流？{LINE}"),
    10: ("", "10", f"Booking 系统与店铺后台是同一个系统，还是各自独立？{LINE}"),
    11: ("", "11", f"送 Takealot 平台仓时，为什么必须为 SKU 贴标？{LINE}"),
    # ── 海运类型 ──
    12: ("海运类型", "1", f"海运到南非的正常周期大概是多久？{LINE}"),
    13: ("", "2", f"【系统操作】在「商品主数据」页面批量导出 CSV 时，文件包含哪些主要字段？{LINE}"),
    14: ("", "3", f"【系统操作】商品或数据的批量导入/导出任务完成后，应去哪个页面查看进度和下载结果？{LINE}"),
    # ── 海外仓类型 ──
    15: ("海外仓类型", "1", f"我们的海外仓位于南非哪个城市？{LINE}"),
    16: ("", "2", f"JHB 仓的客户退货一般什么时候处理？由谁负责跟进？{LINE}"),
    17: ("", "3", f"【系统操作】在「入库实收」页面，WMS 实收数量与预期不一致时，应如何确认和处理？{LINE}"),
    18: ("", "4", f"海外仓团队主要分工有哪些？（如入库、出库、送仓、退货等）{LINE}"),
    19: ("", "5", f"JHB 送 Takealot 仓的费用如何计算？{LINE}"),
    20: ("", "6", f"送 CPT 仓需要提前几天预约？为什么必须提前预约？{LINE}"),
    21: ("", "7", f"DBN 仓的发货/派送形式是什么？{LINE}"),
    22: ("", "8", f"JHB 送仓时外箱破损，通常会导致什么后果？{LINE}"),
    23: ("", "9", f"透明包装且带有 SKU 码的货物，能否送入 Takealot 平台仓？{LINE}"),
    # ── 货盘类型 ──
    24: ("货盘类型", "1", f"货盘模式主要解决了卖家哪些痛点？{LINE}"),
    25: ("", "2", f"【系统操作】新客户首次充值前，在「客户充值」模块必须先完成什么操作？{LINE}"),
    26: ("", "3", f"跨境店目前能否做货盘合作？{LINE}"),
    27: ("", "4", f"货盘货物到达哪个节点后才可以正式开售？{LINE}"),
    28: ("", "5", f"【系统操作】「商品主数据」中商品状态有哪几种？分别表示什么含义？{LINE}"),
    29: ("", "6", f"【系统操作】商品 SKU 批量导入时，模板中哪些字段是必填的？{LINE}"),
    # ── 系统类型（WMS 客户侧 + ERP 内部）──
    30: ("系统类型", "1", f"客户在货物发往南非、对接海外仓 WMS 之前，必须在系统里完成哪项关键操作？{LINE}"),
    31: ("", "2", f"【系统操作】新客户 OMS 开户时，在系统里需要填写/选择哪些信息？{LINE}"),
    32: ("", "3", f"我们目前使用的是哪家公司的 WMS 系统？{LINE}"),
    33: (
        "",
        "4",
        f"请按正确先后顺序排列以下 WMS 操作流程（填写序号 1-7）："
        f"创建出库单、出库换标、下载自定义编码、下载箱唛、创建入库单、创建产品、上传附件。{LINE}",
    ),
    34: ("", "5", f"客户创建出库单送 CPT 仓，除 SKU 换标外，还需要上传哪几个附件？{LINE}"),
    35: ("", "6", f"【系统操作】客户在 ERP 充值的标准操作流程是什么？（从开户到到账）{LINE}"),
    36: ("", "7", f"如果客户使用自定义产品编码出库且未换标就送往 JHB，会产生什么后果？{LINE}"),
    # ── 平台类型 ──
    37: ("平台类型", "1", f"Takealot 三个平台仓分别位于南非哪些城市？{LINE}"),
    38: ("", "2", f"【系统操作】ERP 工作台「全局搜索」可以搜索到哪些类型的业务单据？{LINE}"),
    39: ("", "3", f"Temu、Takealot、Amazon 三个平台在南非的市场占有率，从大到小如何排序？{LINE}"),
    40: ("", "4", f"Takealot 能否实现当日达？哪些情况下可以、哪些情况下不可以？{LINE}"),
    41: ("", "5", f"直邮模式下，货物是送平台仓，还是直接送给消费者？{LINE}"),
    42: ("", "6", f"【系统操作】「库存查询」页面主要用来查看什么信息？{LINE}"),
    43: ("", "7", f"【系统操作】「同步日志」记录了 ERP 与哪些系统之间的接口调用？推送失败后如何处理？{LINE}"),
    44: ("", "8", f"Takealot 平台消费者下单后，多少天内可以申请退货？{LINE}"),
    # ── 系统操作（原市场类型模块，保留 1 道公司业务题）──
    45: ("公司业务", "1", f"特柯洛公司主要业务模块及收入来源包括哪几项？{LINE}"),
    46: ("系统操作", "2", f"入库单从创建到 WMS 确认收货，依次经历哪些状态？{LINE}"),
    47: ("", "3", f"创建入库单时「入仓号」有什么作用？推送 WMS 后能否修改？{LINE}"),
    48: ("", "4", f"入库单草稿保存在哪里？如何继续编辑并正式提交？{LINE}"),
    49: ("", "5", f"通过 CSV 批量导入入库 SKU 时，模板中哪些列是必填的？{LINE}"),
    50: ("", "6", f"客户充值时系统提示「尚未开户」，助理的标准处理步骤是什么？{LINE}"),
    51: ("", "7", f"「客户结算」生成账单后，对客户 OMS 账户余额有什么影响？{LINE}"),
    52: ("", "8", f"「成交管理」中「转 ERP 客户」的作用是什么？与后续 OMS 开户有什么关系？{LINE}"),
    53: ("", "9", f"「异步导出导入」页面可以查看和管理哪些类型的任务？{LINE}"),
    54: ("", "10", f"培训客户使用系统时，演示「创建入库单」的标准步骤应包括哪几个环节？{LINE}"),
    55: ("", "11", f"商品主数据修改后，为什么要关注「同步」状态字段？{LINE}"),
}


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="宋体", size=10, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    if not text:
        return
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, "宋体", 10, bold)


def pick_source() -> str:
    for path in SRC_CANDIDATES:
        if os.path.isfile(path):
            return path
    raise FileNotFoundError("找不到源 Word 文件")


def main():
    src = pick_source()
    doc = Document(src)

    title = doc.paragraphs[0]
    title.text = ""
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Takealot 业务与系统操作考核（一期）")
    set_run_font(run, "黑体", 16, True)
    title.add_run("\n")
    sub_run = title.add_run(
        "适用岗位：系统数据维护 / 客户培训助理  |  含 ERP/WMS 系统操作题  |  考核方式：笔试 + 口试抽查"
    )
    set_run_font(sub_run, "宋体", 10, False)

    table = doc.tables[0]
    headers = ["模块", "序号", "题目", "作答区（题目下方超长横线）", "对错（√/×）", "得分"]
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True)
        set_cell_shading(table.rows[0].cells[j], "D9E2F3")

    for row_idx, (module, seq, question) in QUESTIONS.items():
        row = table.rows[row_idx]
        set_cell_text(row.cells[0], module)
        set_cell_text(row.cells[1], seq)
        set_cell_text(row.cells[2], question)

    doc.save(OUT)
    print(f"Source: {src}")
    print(f"Saved: {OUT}")
    try:
        doc.save(OUT_ORIGINAL)
        print(f"Also updated original: {OUT_ORIGINAL}")
    except PermissionError:
        print("Original file locked — close Word/WeChat to overwrite.")
    system_count = sum(1 for q in QUESTIONS.values() if "【系统操作】" in q[2] or q[0] == "系统操作")
    print(f"Total questions: {len(QUESTIONS)}, system-operation related: {system_count}")


if __name__ == "__main__":
    main()
